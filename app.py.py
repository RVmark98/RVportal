import streamlit as st
import pandas as pd
from docx import Document
from datetime import datetime
import sqlite3
import io

# Setup local database to save quotes for future reference
def init_db():
    conn = sqlite3.connect('rv_quotes.db')
    c = conn.cursor()
    c.execute('''CREATE TABLE IF NOT EXISTS quotes 
                 (ref_no TEXT, client TEXT, date TEXT, total REAL)''')
    conn.commit()
    return conn

conn = init_db()

# Custom Styling
st.set_page_config(page_title="RV Quotation Portal", layout="wide")
st.title("⚡ RV Wind Turbine Spare Parts Portal")

# Manual Details Section
with st.container():
    col1, col2 = st.columns(2)
    with col1:
        ref_no = st.text_input("Reference No.", value="MS/01/54/25-26")
        client_name = st.text_input("Client Name", value="RENEW PRIVATE LIMITED")
    with col2:
        date_str = st.date_input("Quotation Date", datetime.now()).strftime("%d/%m/%Y")
        contact = st.text_input("Kind Attn:", value="Ms. Tannu")

# File Upload for Costing
st.subheader("Step 1: Upload Costing Excel")
uploaded_file = st.file_uploader("Upload Excel with: OEM Code, Description, Qty, Lead Time, Unit Price", type=["xlsx"])

if uploaded_file:
    df = pd.read_excel(uploaded_file)
    
    # Automatic Calculations
    df['Total Price (INR)'] = df['Qty'] * df['Unit Price (INR)']
    subtotal = df['Total Price (INR)'].sum()
    sgst = subtotal * 0.09
    cgst = subtotal * 0.09
    grand_total = subtotal + sgst + cgst

    st.write("### Preview of Quotation Items")
    st.table(df)

    # Financial Summary
    st.write(f"**Subtotal:** ₹{subtotal:,.2f}")
    st.write(f"**SGST (9%):** ₹{sgst:,.2f} | **CGST (9%):** ₹{cgst:,.2f}")
    st.success(f"**Total Value:** ₹{grand_total:,.2f}")

    # Step 2: Generate & Save
    if st.button("Generate & Save Quotation"):
        # Save to database
        c = conn.cursor()
        c.execute("INSERT INTO quotes VALUES (?, ?, ?, ?)", (ref_no, client_name, date_str, grand_total))
        conn.commit()
        
        # Create Word Document
        doc = Document()
        doc.add_paragraph(f"Dated: {date_str}")
        doc.add_paragraph(f"Ref No. {ref_no}")
        doc.add_paragraph(f"\nTO,\n{client_name}\nKind Attn: {contact}")
        
        # Add the parts table
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Table Grid'
        hdrs = table.rows[0].cells
        hdrs[0].text, hdrs[1].text, hdrs[2].text = 'OEM Code', 'Description', 'Qty'
        hdrs[3].text, hdrs[4].text, hdrs[5].text = 'Lead Time', 'Unit Price', 'Total'

        for _, row in df.iterrows():
            cells = table.add_row().cells
            cells[0].text, cells[1].text = str(row['OEM Code']), str(row['Material Description'])
            cells[2].text, cells[3].text = str(row['Qty']), str(row['Lead Time'])
            cells[4].text, cells[5].text = f"{row['Unit Price (INR)']:,}", f"{row['Total Price (INR)']:,}"

        # Footer Terms (Standardized for your business)
        doc.add_paragraph(f"\nTotal Value: {grand_total:,.2f}")
        doc.add_paragraph("\nTerms:")
        doc.add_paragraph("• 50% advance / 50% on dispatch from Chennai Warehouse.")
        doc.add_paragraph("• Validity: 07 days.")
        doc.add_paragraph("• Prices are Ex-mill; transportation extra.")
        
        # Save to buffer
        bio = io.BytesIO()
        doc.save(bio)
        st.download_button(label="Download Quote (Word File)", data=bio.getvalue(), file_name=f"Quote_{ref_no.replace('/','_')}.docx")

# Sidebar for Future Reference
if st.sidebar.button("View Past Quotes"):
    history = pd.read_sql_query("SELECT * FROM quotes", conn)
    st.sidebar.write(history)